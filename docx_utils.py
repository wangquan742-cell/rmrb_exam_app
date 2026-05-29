from io import BytesIO
from zipfile import ZipFile, ZIP_DEFLATED
import re

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Pt, Inches
from docx.oxml.ns import qn


def safe_filename(name: str, max_len: int = 60) -> str:
    name = str(name or "未命名文章")
    name = re.sub(r'[\\/:*?"<>|]+', "_", name)
    name = re.sub(r"\s+", "_", name).strip("_")
    return name[:max_len] or "未命名文章"


def set_doc_style(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Songti SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Songti SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_body_paragraph(doc: Document, text: str):
    if text is None:
        text = ""
    for part in str(text).split("\n"):
        if part.strip():
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(22)
            p.paragraph_format.line_spacing = 1.25
            p.add_run(part.strip())


def add_bullets(doc: Document, items):
    for item in items:
        if item:
            doc.add_paragraph(str(item), style="List Bullet")


def build_rule_based_article_analysis(row: dict) -> dict:
    title = row.get("标题", "")
    tags = str(row.get("主题标签", ""))
    thinking = str(row.get("机关思维", ""))
    summary = row.get("一句话概括", "")
    shenlun = row.get("申论积累方向", "")
    interview = row.get("面试题雏形", "")
    bilingual = row.get("中英表达积累", "")
    preview = row.get("正文预览", "")

    intro = f"这篇文章围绕“{title}”展开，属于{tags or '综合素材'}方向。阅读时可以重点把握文章所体现的政策导向、治理逻辑和规范表达。"
    if preview:
        intro += " 从正文内容看，文章具有较强的政策素材、案例素材和语言表达积累价值。"

    structure = [
        "背景：从人民日报的选题看，文章通常承接国家战略、政策部署或现实治理需要。",
        f"主题：文章核心可归入“{tags}”等方向。",
        f"思维：适合积累“{thinking}”等公务员机关工作思维。",
        "启发：复习时不只看新闻事实，更要提炼可迁移到申论、面试和机关工作表达中的方法论。"
    ]

    shenlun_items = [
        f"申论主题：{tags}",
        f"积累方向：{shenlun}",
        "可用方法：按照背景、问题、原因、对策、意义五段式整理。",
        "写作迁移：可迁移到基层治理、公共服务、高质量发展、作风建设、数字治理等常见主题。"
    ]

    interview_items = [
        f"可转化题型：{row.get('可转化面试题型', '')}",
        f"题目雏形：{interview}",
        "答题思路：先表态，再谈意义或问题，再分析原因，最后提出对策和升华。",
        "开口方式：这类题目既要看到积极意义，也要看到推进过程中可能存在的短板，关键是坚持问题导向和群众导向。"
    ]

    xingce_items = [
        "主旨概括：优先找文章中的政策主线、对策句、总结句。",
        "意图判断：看到“对此、为此、关键是、要、必须”等词，重点关注后面的对策。",
        "逻辑填空：积累人民日报常用搭配，如“系统推进、协同发力、精准施策、久久为功、形成闭环”。",
        "阅读方法：先判断背景段，再找问题意识，最后锁定对策和总结。"
    ]

    return {
        "原文导读": intro,
        "结构拆解": structure,
        "申论积累": shenlun_items,
        "面试积累": interview_items,
        "机关思维": f"本文适合训练：{thinking}。复习时要把新闻材料转化为工作语言，关注群众需求、责任落实、协同治理和闭环整改。",
        "行测言语": xingce_items,
        "中英表达": bilingual,
        "背诵金句": [
            "坚持问题导向，聚焦群众急难愁盼，推动工作从“有没有”向“好不好”转变。",
            "以系统观念统筹推进，以闭环机制压实责任，以群众满意检验成效。",
            "把准方向、找准问题、拿出实招，在服务大局中提升治理效能。"
        ],
        "个人复盘": [
            "这篇文章可以用于哪类申论主题？",
            "如果转化成面试题，我会怎么开口？",
            "其中哪一句规范表达值得背下来？",
            "它体现了哪一种公务员工作思维？"
        ]
    }


def create_article_docx(row: dict, deepseek_analysis: str = "", openai_analysis: str = "") -> bytes:
    doc = Document()
    set_doc_style(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = row.get("标题", "人民日报文章学习笔记")
    add_title(doc, f"人民日报考公学习笔记：{title}")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"日期：{row.get('日期', '')} | 版面：{row.get('版面', '')} | 考公价值分：{row.get('考公价值分', '')}").italic = True

    add_heading(doc, "一、文章基本信息", 1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    info = [
        ("标题", row.get("标题", "")),
        ("日期", row.get("日期", "")),
        ("版面", row.get("版面", "")),
        ("主题标签", row.get("主题标签", "")),
        ("机关思维", row.get("机关思维", "")),
        ("可转化面试题型", row.get("可转化面试题型", "")),
        ("链接", row.get("链接", "")),
    ]
    for k, v in info:
        cells = table.add_row().cells
        cells[0].text = str(k)
        cells[1].text = str(v)

    add_heading(doc, "二、文章原文", 1)
    content = row.get("正文全文", "") or row.get("正文预览", "")
    if content:
        add_body_paragraph(doc, content)
    else:
        add_body_paragraph(doc, "当前表格未保存完整原文。可在下一次抓取时勾选“抓取文章正文”。")

    analysis = build_rule_based_article_analysis(row)

    add_heading(doc, "三、原文导读：不读原文也要掌握什么", 1)
    add_body_paragraph(doc, analysis["原文导读"])

    add_heading(doc, "四、文章结构拆解", 1)
    add_bullets(doc, analysis["结构拆解"])

    add_heading(doc, "五、申论积累", 1)
    add_bullets(doc, analysis["申论积累"])

    add_heading(doc, "六、面试题转化与答题思路", 1)
    add_bullets(doc, analysis["面试积累"])

    add_heading(doc, "七、公务员机关思维积累", 1)
    add_body_paragraph(doc, analysis["机关思维"])

    add_heading(doc, "八、行测言语理解积累", 1)
    add_bullets(doc, analysis["行测言语"])

    add_heading(doc, "九、中英表达积累", 1)
    add_body_paragraph(doc, analysis["中英表达"])

    add_heading(doc, "十、今日背诵金句", 1)
    add_bullets(doc, analysis["背诵金句"])

    if openai_analysis:
        add_heading(doc, "十一、ChatGPT/OpenAI API 深度解析结果", 1)
        add_body_paragraph(doc, openai_analysis)

    if deepseek_analysis:
        add_heading(doc, "十二、DeepSeek API 深度解析结果", 1)
        add_body_paragraph(doc, deepseek_analysis)

    add_heading(doc, "十三、个人复盘区", 1)
    add_bullets(doc, analysis["个人复盘"])
    doc.add_paragraph("\n我的补充笔记：\n\n\n")

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def create_combined_docx(df: pd.DataFrame, task_code: str, top_n: int = 10) -> bytes:
    doc = Document()
    set_doc_style(doc)

    add_title(doc, "人民日报考公考编每日学习笔记")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"任务码：{task_code}").italic = True

    if df.empty:
        add_body_paragraph(doc, "暂无文章。")
    else:
        top_df = df.sort_values(by="考公价值分", ascending=False).head(top_n)
        for idx, (_, row) in enumerate(top_df.iterrows(), start=1):
            if idx > 1:
                doc.add_page_break()
            add_heading(doc, f"第{idx}篇：{row.get('标题', '')}", 1)
            row_dict = row.to_dict()
            analysis = build_rule_based_article_analysis(row_dict)

            add_heading(doc, "1. 原文导读", 2)
            add_body_paragraph(doc, analysis["原文导读"])

            add_heading(doc, "2. 原文内容", 2)
            add_body_paragraph(doc, row_dict.get("正文全文", "") or row_dict.get("正文预览", ""))

            add_heading(doc, "3. 申论积累", 2)
            add_bullets(doc, analysis["申论积累"])

            add_heading(doc, "4. 面试题转化", 2)
            add_bullets(doc, analysis["面试积累"])

            add_heading(doc, "5. 行测言语理解积累", 2)
            add_bullets(doc, analysis["行测言语"])

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def create_article_docx_zip(df: pd.DataFrame, task_code: str, top_n: int = 10, deepseek_df: pd.DataFrame = None, openai_df: pd.DataFrame = None) -> bytes:
    zip_bio = BytesIO()

    deepseek_map = {}
    if deepseek_df is not None and not deepseek_df.empty:
        for _, row in deepseek_df.iterrows():
            deepseek_map[row.get("标题", "")] = row.get("AI解析结果", "")

    openai_map = {}
    if openai_df is not None and not openai_df.empty:
        for _, row in openai_df.iterrows():
            openai_map[row.get("标题", "")] = row.get("AI解析结果", "")

    with ZipFile(zip_bio, "w", ZIP_DEFLATED) as zf:
        if df.empty:
            return zip_bio.getvalue()

        top_df = df.sort_values(by="考公价值分", ascending=False).head(top_n)

        for idx, (_, row) in enumerate(top_df.iterrows(), start=1):
            row_dict = row.to_dict()
            title = row_dict.get("标题", f"文章{idx}")
            deepseek_analysis = deepseek_map.get(title, "")
            openai_analysis = openai_map.get(title, "")
            docx_bytes = create_article_docx(row_dict, deepseek_analysis=deepseek_analysis, openai_analysis=openai_analysis)
            filename = f"{idx:02d}_{safe_filename(title)}.docx"
            zf.writestr(filename, docx_bytes)

    return zip_bio.getvalue()
